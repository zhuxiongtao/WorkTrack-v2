"""
合同解析模块 - 阶段 1+2 重写版

主要改进：
1. 使用 contract_parse task_type（修复「任务模型配置」形同虚设的 bug）
2. PDF 优先用 pdfplumber（保留表格结构）+ 文本<100字自动降级 vision
3. 段落排序按章节加权（标题/期限/价格/违约优先于普通正文）
4. prompt 支持相对日期推算、续约条款识别、阿拉伯-中文混写日期
5. 返回字段级 confidence + source_text（让前端能展示 AI 引用了哪句原文）
6. JSON 解析失败自动重试（lower temperature），原始 response 落日志
"""
import json
import logging
import os
import re
import base64
import struct
from typing import Any
import olefile
from sqlmodel import Session, select
from app.services.ai_service import _get_active_provider, _get_active_provider_full, _get_client
from app.models.model_provider import TaskModelConfig, ModelProvider
from app.exceptions import DocumentParseError
from app.config import settings

logger = logging.getLogger("worktrack.contract")


UPLOAD_DIR = settings.effective_contracts_dir

# ====== 关键章节关键词（用于段落排序加权） ======
HIGH_PRIORITY_KW = [
    # 期限 / 日期
    "期限", "有效期", "履行期", "服务期", "合同期", "生效", "失效", "起始", "截止", "届满", "终止",
    "签订日期", "签署日期", "签约日期", "签订于", "签订于", "签订地点",
    # 价格 / 付款
    "合同金额", "合同价款", "总价", "金额", "价款", "费用", "付款", "支付", "结算",
    # 甲乙方
    "甲方", "乙方", "出卖人", "买受人", "委托方", "受托方", "发包人", "承包人",
    # 重要条款
    "违约", "违约金", "续约", "续签", "自动续", "知识产权", "保密", "验收", "争议", "仲裁", "诉讼", "管辖", "送达", "适用法律", "准据法",
]
# 弱关键词（提升到中等优先级）
MEDIUM_PRIORITY_KW = [
    "本合同", "双方", "约定", "权利", "义务", "应", "须", "不得", "应当", "承诺",
    "日", "月", "年", "%", "万", "元",
]

# 解析模型可接受的最大原文长度（char），超出后按段落优先级截断
MAX_TEXT_CHARS = 16000


# ====================================================================
# 供应商 / 视觉模型解析
# ====================================================================
def _resolve_vision_provider(db: Session, user_id: int):
    vision_cfg = db.exec(
        select(TaskModelConfig).where(
            TaskModelConfig.task_type == "vision",
            TaskModelConfig.user_id == user_id,
        )
    ).first()
    if not vision_cfg:
        vision_cfg = db.exec(
            select(TaskModelConfig).where(
                TaskModelConfig.task_type == "vision",
                TaskModelConfig.user_id == None,
            )
        ).first()
    if not vision_cfg or not vision_cfg.provider_id or not vision_cfg.model_name:
        raise RuntimeError("未配置视觉模型")
    provider = db.get(ModelProvider, vision_cfg.provider_id)
    if not provider or not provider.is_active or not provider.api_key:
        raise RuntimeError("视觉模型供应商未激活")
    return provider, vision_cfg.model_name


def _extract_text_via_vision(file_path: str, db: Session, user_id: int = 0) -> str:
    """逐页用 vision 模型 OCR（fallback 路径：仅在 PDF/docx 文本提取为空时调用）"""
    import fitz
    provider, model_name = _resolve_vision_provider(db, user_id)
    base_url, api_key, _, _ = _get_active_provider(db, "vision", user_id)
    client = _get_client(base_url, api_key, provider)

    doc = fitz.open(file_path)
    all_text = ""
    for page_num in range(len(doc)):
        pix = doc[page_num].get_pixmap(dpi=150)
        img_b64 = base64.b64encode(pix.tobytes("png")).decode("utf-8")
        response = client.chat.completions.create(
            model=model_name,
            messages=[{
                "role": "user",
                "content": [
                    {"type": "text", "text": "请逐字识别并提取图片中所有文字，直接输出文字内容，保留段落结构。"},
                    {"type": "image_url", "image_url": {"url": f"data:image/png;base64,{img_b64}"}}
                ]
            }],
            temperature=0.1,
        )
        all_text += (response.choices[0].message.content or "") + "\n\n"
    doc.close()
    return all_text.strip()


# ====================================================================
# 文件文本提取
# ====================================================================
def extract_text_from_pdf(file_path: str) -> str:
    """
    PDF 文本提取（优先 pdfplumber 保表格，回退 fitz）
    返回空字符串表示 PDF 是扫描件或解析失败，需走 vision
    """
    # 优先 pdfplumber（保表格结构）
    try:
        import pdfplumber
        text_parts = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                # 提取文本（保留顺序）
                t = page.extract_text() or ""
                # 尝试提取表格
                try:
                    tables = page.extract_tables() or []
                    for tbl in tables:
                        rows = [" | ".join(str(c) if c is not None else "" for c in row) for row in tbl if row]
                        if rows:
                            t += "\n[表格]\n" + "\n".join(rows) + "\n[/表格]\n"
                except Exception:
                    pass
                if t.strip():
                    text_parts.append(t)
        full = "\n".join(text_parts).strip()
        if len(full) > 50:
            return full
    except ImportError:
        pass
    except Exception as e:
        logger.warning("pdfplumber 解析失败, 回退 fitz: %s", e)

    # 回退 fitz
    try:
        import fitz
        doc = fitz.open(file_path)
        text = ""
        for page in doc:
            text += page.get_text() + "\n"
        doc.close()
        return text.strip()
    except Exception as e:
        logger.error("fitz 解析 PDF 失败: %s", e)
        return ""


def extract_text_from_docx(file_path: str) -> str:
    try:
        from docx import Document
        doc = Document(file_path)
        parts = []
        for p in doc.paragraphs:
            if p.text.strip():
                parts.append(p.text)
        # 提取表格
        try:
            for tbl in doc.tables:
                for row in tbl.rows:
                    cells = [c.text.strip() for c in row.cells]
                    if any(cells):
                        parts.append(" | ".join(cells))
        except Exception:
            pass
        return "\n".join(parts)
    except Exception as e:
        logger.error("提取 docx 文本失败: %s", e)
        return ""


def _extract_text_from_doc_via_olefile(file_path: str) -> str:
    """
    纯 Python 解码老格式 .doc（OLE2 / WPS / 旧版 Word 生成的 .doc）。
    策略：olefile 读取 WordDocument 流，按 FIB.fcMin/fcMac 截取正文，按 UTF-16 LE 解码。
    适用于：WPS Office、LibreOffice 7+、MS Word 2010+ 保存的 .doc（未压缩格式）。
    不适用于：使用 Word 97/2000 压缩算法存储的 .doc（会回退到空串，让上层走 vision 兜底）。
    """
    try:
        ole = olefile.OleFileIO(file_path)
    except Exception as e:
        logger.debug("olefile 打开失败: %s", e)
        return ""

    if not ole.exists("WordDocument"):
        ole.close()
        return ""

    try:
        wd = ole.openstream("WordDocument").read()
    except Exception as e:
        logger.debug("读取 WordDocument 流失败: %s", e)
        ole.close()
        return ""

    ole.close()

    if len(wd) < 0x20:
        return ""

    # 验证 FIB 头 magic = 0xA5EC（小端 = ec a5）
    if wd[0] != 0xEC or wd[1] != 0xA5:
        return ""

    # fcMin/fcMac：正文在 WordDocument 流中的字节偏移
    fcMin = struct.unpack("<I", wd[0x18:0x1C])[0]
    fcMac = struct.unpack("<I", wd[0x1C:0x20])[0]

    if fcMin < 0 or fcMac <= fcMin or fcMac > len(wd):
        return ""

    text_bytes = wd[fcMin:fcMac]

    # 启发式判断：未压缩文档正文是 UTF-16 LE
    # 先试着按 UTF-16 LE 解码，过滤控制字符和乱码后看是否含有中文/可读 ASCII
    try:
        text = text_bytes.decode("utf-16-le", errors="ignore")
    except Exception:
        return ""

    # 过滤控制字符（除 \n \r \t 之外）
    cleaned_chars = []
    for c in text:
        if c in "\n\r\t":
            cleaned_chars.append(c)
        elif c.isprintable() or ('\u4e00' <= c <= '\u9fff'):
            cleaned_chars.append(c)
        elif c == "\x00":
            continue
        else:
            cleaned_chars.append(" ")
    cleaned = "".join(cleaned_chars)

    # 折叠多余空白
    cleaned = re.sub(r"[ \t]+", " ", cleaned)
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
    cleaned = "\n".join(line.strip() for line in cleaned.split("\n") if line.strip())

    # 验证：至少有 30 个中文字符 或 200 个可读字符 才认为成功
    chinese_count = sum(1 for c in cleaned if '\u4e00' <= c <= '\u9fff')
    if chinese_count < 30 and len(cleaned) < 200:
        return ""

    return cleaned


def extract_text_from_legacy_doc(file_path: str) -> str:
    """老格式 .doc 解析（多策略 fallback）"""
    # 策略 0（新增）：纯 Python 用 olefile 直接解码 UTF-16 LE 正文
    # 覆盖 WPS、LibreOffice 7+、MS Word 2010+ 保存的未压缩 .doc
    try:
        text = _extract_text_from_doc_via_olefile(file_path)
        if text and len(text) > 50:
            return text
    except Exception as e:
        logger.debug("olefile 策略失败: %s", e)

    # 策略 1：先试 docx 库（少数 .doc 文件实际是 docx 格式）
    try:
        from docx import Document
        doc = Document(file_path)
        text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        if len(text) > 50:
            return text
    except Exception:
        pass

    # 策略 2：antiword（Linux/Mac 命令行，Windows 通常无）
    import tempfile, subprocess, os as _os, platform
    tmp_txt = tempfile.NamedTemporaryFile(suffix=".txt", delete=False).name
    try:
        if platform.system() in ("Linux", "Darwin"):
            try:
                r = subprocess.run(
                    ["antiword", "-t", file_path],
                    capture_output=True, text=True, timeout=30
                )
                if r.returncode == 0 and r.stdout and len(r.stdout) > 50:
                    return r.stdout.strip()
            except (FileNotFoundError, subprocess.TimeoutExpired):
                pass

        # 策略 3：libreoffice headless 转 docx 再解析（Linux/Docker 友好）
        if platform.system() == "Linux":
            try:
                tmp_dir = tempfile.mkdtemp()
                r = subprocess.run(
                    ["libreoffice", "--headless", "--convert-to", "docx", "--outdir", tmp_dir, file_path],
                    capture_output=True, text=True, timeout=60
                )
                if r.returncode == 0:
                    converted = _os.path.join(tmp_dir, _os.path.splitext(_os.path.basename(file_path))[0] + ".docx")
                    if _os.path.exists(converted):
                        result = extract_text_from_docx(converted)
                        if result:
                            return result
            except (FileNotFoundError, subprocess.TimeoutExpired):
                pass
            except Exception as e:
                logger.warning("libreoffice 转换失败: %s", e)

        # 策略 4：Windows win32com（仅当装了 MS Word）
        if platform.system() == "Windows":
            try:
                import win32com.client
                word = win32com.client.Dispatch("Word.Application")
                word.Visible = False
                word.DisplayAlerts = 0
                doc = None
                try:
                    doc = word.Documents.Open(file_path)
                    tmp_pdf = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False).name
                    doc.SaveAs(tmp_pdf, FileFormat=17)
                finally:
                    if doc:
                        try: doc.Close()
                        except Exception: pass
                    try: word.Quit()
                    except Exception: pass

                if _os.path.exists(tmp_pdf) and _os.path.getsize(tmp_pdf) > 100:
                    result = extract_text_from_pdf(tmp_pdf)
                    try: _os.remove(tmp_pdf)
                    except Exception: pass
                    if result:
                        return result
            except Exception as e:
                logger.warning("Windows win32com 转换失败: %s", e)

        return ""
    except Exception as e:
        logger.error("legacy doc 转换失败: %s", e)
        return ""
    finally:
        try:
            _os.remove(tmp_txt)
        except Exception:
            pass


def extract_text(file_path: str, file_type: str) -> str:
    if file_type in (".pdf", "pdf"):
        return extract_text_from_pdf(file_path)
    elif file_type in (".docx", "docx"):
        return extract_text_from_docx(file_path)
    elif file_type in (".doc", "doc"):
        return extract_text_from_legacy_doc(file_path)
    return ""


def extract_text_with_vision_fallback(file_path: str, file_type: str, db: Session, user_id: int = 0) -> str:
    """
    文本提取 + 自动 OCR fallback
    返回：提取的文本（不会抛错；失败时返回空字符串或部分文本）
    """
    text = extract_text(file_path, file_type)
    # 文本太短（< 100 字）很可能是扫描件，自动降级 vision
    if len(text) < 100 and file_type in (".pdf", "pdf", ".docx", "docx", ".doc", "doc"):
        try:
            text = _extract_text_via_vision(file_path, db, user_id)
        except RuntimeError as e:
            # vision 也没配 - 抛出明确错误（带可执行的解决建议）
            raise RuntimeError(
                f"此文件无法直接提取文字（可能为扫描件 / 老格式 .doc 无 MS Word 解码 / 含图片型 PDF），"
                f"需要配置视觉模型来识别：{e}。"
                f"建议：① 在「AI 模型 → 任务模型配置」中添加 vision 任务模型；"
                f"或 ② Windows 安装 MS Word 以解码 .doc；"
                f"或 ③ 把文件另存为 .docx / 纯文本 PDF 后重新上传。"
            )
        except Exception as e:
            logger.error("vision 识别失败: %s", e)
            # 不抛错，让上层决定如何处理
    return text


# ====================================================================
# 关键段落排序
# ====================================================================
def _extract_key_sections(raw_text: str) -> str:
    """
    从合同原文中提取包含关键信息的段落，按优先级排序，超长时保留高优先级段
    """
    if len(raw_text) <= MAX_TEXT_CHARS:
        return raw_text

    lines = raw_text.split("\n")
    scored = []
    for line in lines:
        s = line.strip()
        if not s:
            continue
        score = 0
        for kw in HIGH_PRIORITY_KW:
            if kw in s:
                score += 2
        for kw in MEDIUM_PRIORITY_KW:
            if kw in s:
                score += 1
        scored.append((score, s))

    # 按分数倒序，再按原顺序
    scored.sort(key=lambda x: -x[0])

    result_parts = []
    total = 0
    for score, s in scored:
        if total + len(s) + 1 > MAX_TEXT_CHARS:
            if score < 2:
                # 低优先级直接砍
                continue
            else:
                # 高优先级截断保留
                remain = MAX_TEXT_CHARS - total
                if remain > 100:
                    result_parts.append(s[:remain])
                    total = MAX_TEXT_CHARS
                break
        result_parts.append(s)
        total += len(s) + 1

    return "\n".join(result_parts)


# ====================================================================
# 解析 prompt 与核心解析函数
# ====================================================================
_PARSE_PROMPT = """你是一个专业的合同分析助手。请仔细分析以下合同文本，提取所有关键字段并以严格 JSON 格式返回。

## 必须返回的字段（每个字段都包含 value、confidence 0-1、source_text 原文引用）

### 基本信息
- `title`: 合同名称（如"采购合同"、"技术服务协议"等）
- `contract_type`: 合同类型枚举之一：销售合同 / 采购合同 / 服务合同 / 租赁合同 / 劳动合同 / 保密协议 / 技术合同 / 咨询合同 / 合作协议 / 其他
- `contract_no`: 合同编号
- `sign_date`: 签订日期 (YYYY-MM-DD)
- `sign_location`: 签订地点

### 合同主体
- `party_a`: 甲方全称
- `party_b`: 乙方全称

### 日期与期限（最重要！）
- `start_date`: 合同生效/开始日期 (YYYY-MM-DD)
- `end_date`: 合同失效/截止/到期日期 (YYYY-MM-DD)
- `effective_term`: 合同期限的原文描述（如果推算不出具体日期，保留原文如"自签订之日起 3 年"）

#### 日期识别完整规则
**格式覆盖**：
- 阿拉伯数字："2024年1月1日"、"2024.1.1"、"2024/1/1"、"2024-1-1"
- 阿拉伯-中文混写："二〇二四年一月一日"、"二零二四年一月一日"（务必转换）
- 农历/财政年度：如明确说明是农历或财年，请注明

**相对日期推算**（极其重要）：
- "自签订日起 3 年" → 若 sign_date 已知：end_date = sign_date + 3年
- "自生效日起 36 个月" → end_date = start_date + 36个月
- "服务期限 1 年" + "自 X 日起" → 推算
- **如果推算结果填入 end_date，同时在 effective_term 保留原文**

**续约条款识别**（独立字段）：
- 看到"自动续约"、"续签"、"期满"、"到期前 30 日"、"未提出书面异议"等表述时
- 将完整条款原文提取到 `auto_renew` 字段
- 若有续约，end_date 仍按原合同期推算，在 auto_renew 字段说明续约规则

**特别注意**：
- end_date 是合同**终止/失效/到期**日期，**不是**签订日期
- "本合同有效期 3 年" + "签订于 2024.1.1" → 推算 end_date 为 2027.1.1
- 若 start_date 推不出而 end_date 推出，start_date 填 null 但 effective_term 保留原文

### 金额
- `contract_amount`: 合同总金额（数字，单位：万元）
- `currency`: 币种（CNY/USD/EUR 等）
- `amount_in_words`: 金额大写（"人民币捌拾伍万捌仟伍佰元整"）
- `payment_schedule`: 付款节点列表 JSON 字符串，格式如：
  `[{{"phase":"预付款","percent":30,"condition":"合同签订后 5 日内"}},{{"phase":"验收款","percent":60,"condition":"验收合格后"}},{{"phase":"质保金","percent":10,"condition":"质保期满"}}]`
  - 若无明确节点，原文描述如"分期支付：30%/60%/10%"
- `payment_terms`: 付款方式文字描述（电汇/承兑/信用证等）

### 关键条款（每条都要原文化）
- `penalty_clause`: 违约金条款（比例/计算方式/上限）
- `acceptance_terms`: 验收条款（验收标准、验收期限）
- `ip_clause`: 知识产权归属
- `dispute_resolution`: 争议解决（仲裁机构/法院管辖地）
- `governing_law`: 适用法律
- `notice_clause`: 通知与送达条款
- `confidentiality`: 保密条款
- `key_clauses`: 其他重要条款（合并摘要）

### 其他
- `summary`: 合同核心内容 2-3 句话摘要
- `extraction_notes`: 解析备注（哪些字段不确定、哪些是推算的）

## 字段填写规范
- 每个字段用对象表示：`{{"value": ..., "confidence": 0.0-1.0, "source_text": "原文引用"}}`
- 拿不准的字段：value 仍要尽力提取，confidence 调低（0.3-0.6）
- 完全找不到：value 填 null，confidence 为 0
- **source_text 必须是合同中实际存在的句子或短语**，不要编造

## 输出格式（纯 JSON，无 markdown 代码块，无其他文字）
```json
{{
  "title": {{"value": "...", "confidence": 0.9, "source_text": "..."}},
  "contract_type": {{"value": "...", "confidence": 0.85, "source_text": "..."}},
  "contract_no": {{"value": "...", "confidence": 0.95, "source_text": "..."}},
  "sign_date": {{"value": "YYYY-MM-DD", "confidence": 0.9, "source_text": "..."}},
  "sign_location": {{"value": "...", "confidence": 0.8, "source_text": "..."}},
  "party_a": {{"value": "...", "confidence": 0.95, "source_text": "..."}},
  "party_b": {{"value": "...", "confidence": 0.95, "source_text": "..."}},
  "start_date": {{"value": "YYYY-MM-DD", "confidence": 0.85, "source_text": "..."}},
  "end_date": {{"value": "YYYY-MM-DD", "confidence": 0.8, "source_text": "..."}},
  "effective_term": {{"value": "...", "confidence": 0.9, "source_text": "..."}},
  "auto_renew": {{"value": "...", "confidence": 0.7, "source_text": "..."}},
  "contract_amount": {{"value": 85.85, "confidence": 0.95, "source_text": "..."}},
  "currency": {{"value": "CNY", "confidence": 0.95, "source_text": "..."}},
  "amount_in_words": {{"value": "...", "confidence": 0.9, "source_text": "..."}},
  "payment_schedule": {{"value": "[{{...}}]", "confidence": 0.7, "source_text": "..."}},
  "payment_terms": {{"value": "...", "confidence": 0.85, "source_text": "..."}},
  "penalty_clause": {{"value": "...", "confidence": 0.85, "source_text": "..."}},
  "acceptance_terms": {{"value": "...", "confidence": 0.85, "source_text": "..."}},
  "ip_clause": {{"value": "...", "confidence": 0.85, "source_text": "..."}},
  "dispute_resolution": {{"value": "...", "confidence": 0.9, "source_text": "..."}},
  "governing_law": {{"value": "...", "confidence": 0.9, "source_text": "..."}},
  "notice_clause": {{"value": "...", "confidence": 0.8, "source_text": "..."}},
  "confidentiality": {{"value": "...", "confidence": 0.8, "source_text": "..."}},
  "key_clauses": {{"value": "...", "confidence": 0.85, "source_text": "..."}},
  "summary": {{"value": "...", "confidence": 0.9, "source_text": "..."}},
  "extraction_notes": {{"value": "...", "confidence": 1.0, "source_text": ""}}
}}
```

合同文本：
{focused_text}
"""


def _normalize_field(field: Any) -> dict:
    """规范化字段输出，确保一定有 value/confidence/source_text 三个键"""
    if field is None:
        return {"value": None, "confidence": 0.0, "source_text": ""}
    if isinstance(field, dict):
        return {
            "value": field.get("value"),
            "confidence": float(field.get("confidence") or 0),
            "source_text": field.get("source_text") or "",
        }
    # 兼容旧格式：直接给值（说明是旧版 prompt 输出）
    return {"value": field, "confidence": 0.5, "source_text": ""}


def _safe_get_value(parsed: dict, key: str) -> Any:
    """从解析结果中安全取值，兼容新旧两种 prompt 输出格式"""
    raw = parsed.get(key)
    if isinstance(raw, dict) and "value" in raw:
        return raw.get("value")
    return raw


def _safe_get_confidence(parsed: dict, key: str) -> float:
    raw = parsed.get(key)
    if isinstance(raw, dict):
        try:
            return float(raw.get("confidence") or 0)
        except (TypeError, ValueError):
            return 0.0
    return 0.5  # 旧格式无 confidence 时给 0.5


def _safe_get_source(parsed: dict, key: str) -> str:
    raw = parsed.get(key)
    if isinstance(raw, dict):
        return str(raw.get("source_text") or "")
    return ""


def _try_parse_json(text: str) -> dict:
    """尝试从 LLM 响应中解析 JSON，多种容错策略"""
    text = text.strip()
    # 1. 直接解析
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        pass
    # 2. 去掉 markdown 代码块
    if text.startswith("```"):
        lines = text.split("\n")
        text = "\n".join(lines[1:])
        if text.endswith("```"):
            text = text[:-3]
        text = text.strip()
        try:
            return json.loads(text)
        except json.JSONDecodeError:
            pass
    # 3. 提取第一个完整的 { ... } 块
    m = re.search(r'\{[\s\S]*\}', text)
    if m:
        try:
            return json.loads(m.group())
        except json.JSONDecodeError:
            pass
    raise ValueError(f"无法从 LLM 响应中解析 JSON: {text[:200]}")


def parse_contract(raw_text: str, db: Session, user_id: int = 0) -> dict:
    """
    合同解析主入口
    关键改进：
    - 使用 contract_parse task_type（用户可在「任务模型配置」专门配置）
    - JSON 解析失败自动重试 1 次（lower temperature）
    - 返回的 result 包含 _raw_response / _parse_error（用于诊断）
    """
    # 修复 Bug #1：使用 contract_parse task type
    try:
        base_url, api_key, model, provider, task_cfg, pm = _get_active_provider_full(db, "contract_parse", user_id)
    except Exception as e:
        # 兜底：没配 contract_parse 就用 chat
        logger.warning("未配置 contract_parse 任务模型, 降级到 chat: %s", e)
        base_url, api_key, model, provider, task_cfg, pm = _get_active_provider_full(db, "chat", user_id)
    client = _get_client(base_url, api_key, provider)

    focused_text = _extract_key_sections(raw_text)

    # 预先准备 2 套参数：① 带 response_format（标准 OpenAI 模型）② 不带（国产/小模型如 MiMo 不支持）
    from app.services.ai_service import resolve_chat_params
    params_with_rf = resolve_chat_params(
        db, model=pm, task_cfg=task_cfg,
        func_defaults={"temperature": 0.1},
        func_overrides={},
    )
    # 构造一份去掉 response_format 的副本
    params_no_rf = {k: v for k, v in params_with_rf.items() if k != "response_format"}

    last_error = None
    last_response = None
    # (temp, suffix, params) 组合：第一次带 response_format，第二次去掉
    for attempt, (temp, suffix, params) in enumerate([
        (0.1, "", params_with_rf),
        (0.0, "\n\n注意：上次输出不是有效 JSON，请严格只返回纯 JSON 对象。", params_with_rf),
        (0.1, "", params_no_rf),  # 兜底：去掉 response_format（兼容 MiMo 等小模型）
        (0.0, "\n\n注意：上次输出不是有效 JSON，请严格只返回纯 JSON 对象。", params_no_rf),
    ]):
        try:
            response = client.chat.completions.create(
                model=model,
                messages=[
                    {"role": "system", "content": "你是一个专业的合同分析助手，擅长从中文合同中精准抽取日期、金额、甲乙方、关键条款。请严格按照 JSON 格式输出，不要添加任何解释文字。"},
                    {"role": "user", "content": _PARSE_PROMPT.format(focused_text=focused_text[:MAX_TEXT_CHARS]) + suffix},
                ],
                **params,
            )
            content = (response.choices[0].message.content or "").strip()
            last_response = content
            parsed = _try_parse_json(content)
            # 成功解析
            parsed["_raw_response"] = content[:2000]  # 最多保留 2K 用于诊断
            parsed["_attempts"] = attempt + 1
            return parsed
        except Exception as e:
            last_error = e
            logger.warning("合同解析第 %d 次失败: %s", attempt + 1, e)
            continue

    # 4 次都失败
    err_msg = str(last_error)[:300] if last_error else "未知错误"
    # 把 "phase" 这种不知所谓的错误翻译成可读提示
    if err_msg.strip() in ('"phase"', 'phase') or len(err_msg) < 20 and "phase" in err_msg:
        err_msg = (
            f"AI 模型调用被服务端拒绝（{err_msg}）。"
            f"这通常表示：① 当前模型（如 MiMo/小模型）不支持 response_format=json_object 参数；"
            f"或 ② 模型供应商 API key / endpoint 配置异常。"
            f"建议：换用支持 JSON mode 的模型（如 deepseek、qwen、gpt 系列），"
            f"或在「AI 模型 → 任务模型配置」中把该模型的 response_format 改为 text。"
        )
    return {
        "error": err_msg,
        "_raw_response": (last_response or "")[:2000],
        "_attempts": 4,
    }


def apply_parse_result(contract, parsed: dict) -> None:
    """
    把 parse_contract 的结果应用到 contract 对象上
    自动写入解析状态、错误信息、extraction_meta 等
    """
    from datetime import datetime, timezone

    if "error" in parsed:
        contract.parse_status = "failed"
        contract.parse_error = parsed.get("error", "")[:500]
        contract.parsed_at = datetime.now(timezone.utc)
        return

    try:
        # 字符串/日期字段
        if _safe_get_value(parsed, "sign_date"):
            try:
                from datetime import date
                v = _safe_get_value(parsed, "sign_date")
                contract.sign_date = date.fromisoformat(str(v))
            except Exception:
                pass
        if _safe_get_value(parsed, "start_date"):
            try:
                from datetime import date
                v = _safe_get_value(parsed, "start_date")
                contract.start_date = date.fromisoformat(str(v))
            except Exception:
                pass
        if _safe_get_value(parsed, "end_date"):
            try:
                from datetime import date
                v = _safe_get_value(parsed, "end_date")
                contract.end_date = date.fromisoformat(str(v))
            except Exception:
                pass

        # 字符串字段
        s = lambda k: str(_safe_get_value(parsed, k) or "").strip()
        contract.party_a = s("party_a") or contract.party_a
        contract.party_b = s("party_b") or contract.party_b
        contract.contract_type = s("contract_type")
        contract.effective_term = s("effective_term")
        contract.auto_renew = s("auto_renew")
        contract.penalty_clause = s("penalty_clause")
        contract.acceptance_terms = s("acceptance_terms")
        contract.ip_clause = s("ip_clause")
        contract.dispute_resolution = s("dispute_resolution")
        contract.governing_law = s("governing_law")
        contract.notice_clause = s("notice_clause")
        contract.summary = s("summary")
        contract.key_clauses = s("key_clauses")
        contract.payment_terms = s("payment_terms")

        # 付款节点（JSON 字符串）
        ps = _safe_get_value(parsed, "payment_schedule")
        if ps:
            # 验证是否为合法 JSON
            if isinstance(ps, str):
                try:
                    json.loads(ps)
                    contract.payment_schedule = ps
                except Exception:
                    contract.payment_schedule = json.dumps([{"phase": "付款", "condition": ps}], ensure_ascii=False)
            elif isinstance(ps, list):
                contract.payment_schedule = json.dumps(ps, ensure_ascii=False)

        # 金额
        amt = _safe_get_value(parsed, "contract_amount")
        if amt is not None:
            try:
                contract.contract_amount = float(amt)
            except (TypeError, ValueError):
                pass
        cur = s("currency")
        if cur:
            contract.currency = cur

        # 抽取元数据
        meta = {}
        for key in ["title", "contract_type", "contract_no", "sign_date", "start_date", "end_date",
                    "party_a", "party_b", "contract_amount", "payment_schedule",
                    "penalty_clause", "acceptance_terms", "ip_clause", "dispute_resolution",
                    "governing_law", "auto_renew", "effective_term"]:
            meta[key] = {
                "confidence": _safe_get_confidence(parsed, key),
                "source_text": _safe_get_source(parsed, key),
            }
        contract.extraction_meta = json.dumps(meta, ensure_ascii=False)

        contract.parse_status = "success"
        contract.parse_error = ""
        contract.parsed_at = datetime.now(timezone.utc)
    except Exception as e:
        logger.exception("apply_parse_result 失败: %s", e)
        contract.parse_status = "failed"
        contract.parse_error = f"结果回填失败: {str(e)[:300]}"
        contract.parsed_at = datetime.now(timezone.utc)
